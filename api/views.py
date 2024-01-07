import os

from django.shortcuts import render
from django.http import HttpResponse, JsonResponse, HttpResponseRedirect
from django.forms import model_to_dict

import time
import json
import uuid

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from api.models import BidInfo, ProjectList

from shibing624.shibing624 import get_project


# from shibing624.add_data import add, data
# add(data)


def get_select_bid_info(request):
    # 将每一行的project_name,project_detail,budget,start_time, buyer_info进行拼接(str_bid_info)，
    # 使用shibing64中文模型将接受到的select_bid_info和拼接后的字符串(str_bid_info)进行相似度计算，
    # 找出相似度最高的前五个。并传回：[{1.:推荐一},{2.:推荐二},{3.:推荐3}...]

    # if request.method == "GET":
    if request.method == "POST":
        # print(request.POST)
        select_bid_info = request.POST.get('select_bid_info')

        projects = BidInfo.objects.filter().all()

        projects = get_project(select_bid_info, projects)
        result = [model_to_dict(project) for project in projects]

        # print(result)

        code = 1
        msg = "ok"
        data = dict(code=code, msg=msg, result=result)
        return JsonResponse(data, json_dumps_params={'ensure_ascii': False})


def get_docx_info(docx_path):
    # 打开原始docx文件
    doc = Document(docx_path)

    docx_info = {}
    title = "title1"

    # 遍历所有段落
    for paragraph in doc.paragraphs:
        paragraph_style = paragraph.style.name

        if paragraph_style.startswith("Title"):
            title = paragraph.text
            if title not in docx_info:
                docx_info[title] = []
        else:
            docx_info[title].append(paragraph.text)

    return docx_info

    # # 遍历文档中的段落
    # for paragraph in doc.paragraphs:
    #     # 获取段落的样式
    #     paragraph_style = paragraph.style.name
    #     print(f'Style: {paragraph_style} Paragraph: {paragraph.text}')

    # # 遍历文档中的表格
    # for table in doc.tables:
    #     # 遍历表格中的行
    #     for row in table.rows:
    #         # 遍历行中的单元格
    #         for cell in row.cells:
    #             # 获取单元格中的段落
    #             for paragraph in cell.paragraphs:
    #                 # 获取段落的样式
    #                 paragraph_style = paragraph.style.name
    #                 print(f'Style: {paragraph_style} Cell Text: {paragraph.text}')


def get_tech_docx_view(request):
    # 1. get_tech_docx_view(project_number): 请求接口路径，project_number 项目编     号
    # docx_info:返回值为 null或者 {title1:context,title2:context,title3:context.....} context为字符串列表.
    # 描述：
    # 创建 get_tech_docx_view(project_number)接口，通过参数(project_number)查询数据库 project_list ，不存在直接返回 null。
    # 否则将查询到的对应的标书地址获取，并通过地址读取标书，将文档的标题和内容返回，其中标题作为key，对应标题下的所有内容作为value。
    # 返回参数为: docx_info{title1:context, title2:context, title3:context...}, context 为字符串列表（存放对应标题下的所有段落，每个段落作为数组中的一项）。

    # if request.method == "GET":
    if request.method == "POST":
        # print(request.POST)
        project_number = request.POST.get('project_number')
        project_name = request.POST.get('project_name')
        try:
            project = ProjectList.objects.filter(project_number=project_number, project_name=project_name).first()
        except Exception as e:
            print(e)

        if project:
            docx_path = f"{project.project_link}"
            docx_info = get_docx_info(docx_path)

            code = 1
            msg = "ok"
            data = dict(code=code, msg=msg, docx_info=docx_info)
        else:
            code = 0
            msg = "null"
            data = dict(code=code, msg=msg, docx_info=None)

        return JsonResponse(data, json_dumps_params={'ensure_ascii': False})


def get_add_info_p_view(request):
    # 2. get_add_info_p_view(add_info_p): 请求接口url地址。
    # add_info_p:前端传来的参数，为json数组：{‘title’:title,’context’:context,’input_info_p’:input_info_p}
    # title: 标题，context：内容，input_info_p：用户输入的信息
    # 描述：
    # 创建get_add_info_p_view(add_info_p) 接口，将传入的参数，解析add_info_p数据，
    # 将标题和内容以及用户输入的信息按照标题(参数中的title)，
    # 内容(参数中的context, （context和input_info_p中间加换行）input_info_p,)的格式排版保存为docx文件。

    # if request.method == "GET":
    if request.method == "POST":
        # print(request.POST)
        title = request.POST.get('title')
        context = request.POST.get('context')
        input_info_p = request.POST.get('input_info_p')
        # json_data = json.loads(add_info_p)

        # json_data = {'title': '案件资料', 'context': '根据比选文件的要求，结合本项目的特点，从项目实施组织及部署的科学性；宣传氛围布置工序安排的合理性；具体布置方法选用的技术性、经济性和实现的可能性进行了科学的论断和详细的阐述。", "针对一些技术难点提出了解决问题的方法。对关键工序进行了合理的编排。', 'input_info_p': '对工期、质量、安全、文明宣传布置提出了工作目标，并针对各项目建立了保证体系，制定了相应的技术保证措施。", "从采购单位利益及项目顺利进行的角度上考虑制定了与采购单位等其他合作单位的配合措施。'}

        doc = Document()

        # 添加标题
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True  # 设置标题为粗体
        title_run.font.size = Pt(16)  # 设置标题字体大小
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置标题居中

        doc.add_paragraph(context)
        doc.add_paragraph("")
        doc.add_paragraph(input_info_p)

        doc.save(
            f"media/project_list_add/{time.strftime('%Y_%m_%d_%H_%M_%S', time.localtime())}_{uuid.uuid4().hex}.docx")

        code = 1
        msg = "ok"
        data = dict(code=code, msg=msg,
                    file_path=f"media/project_list_add/{time.strftime('%Y_%m_%d_%H_%M_%S', time.localtime())}_{uuid.uuid4().hex}.docx")
        return JsonResponse(data, json_dumps_params={'ensure_ascii': False})


def save_file_data(file, name, save_directory):
    new_file_name = f'{name}.{file.name.split(".")[-1]}'
    _file_path = os.path.join(save_directory, new_file_name)
    with open(_file_path, 'wb') as destination:
        for chunk in file.chunks():
            destination.write(chunk)
    file.seek(0)  # 将文件指针移动到文件开头
    return new_file_name


def record_post(request):
    # 后端：创建 record_post 接口，接收前台传来的音频信息，return值为wav或者mp3文件，不向前端返回值，作为其它模块的一个调用函数。

    # if request.method == "GET":
    if request.method == "POST":
        # print(request.POST)
        record_file = request.FILES.get("record_file")

        save_directory = f"media/record_file/"
        # os.makedirs(save_directory, exist_ok=True)  # 创建目录

        if record_file:
            file_name = save_file_data(record_file,
                                       f"{time.strftime('%Y_%m_%d_%H_%M_%S', time.localtime())}_{uuid.uuid4().hex}",
                                       save_directory)
        else:
            file_name = "null"

        code = 1
        msg = "ok"
        data = dict(code=code, msg=msg, file_path=f"media/record_file/{file_name}")
        return JsonResponse(data, json_dumps_params={'ensure_ascii': False})
